from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/magnus/Documents/Github/fruitfly-monitoring")
INPUT_PATH = ROOT / "docs" / "fruitfly_technical_report.md"
OUTPUT_PATH = ROOT / "docs" / "fruitfly_technical_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10.5)
        else:
            paragraph.add_run(part)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_cover(document: Document) -> None:
    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = 1
    title_run = title.add_run("Technical Report on the Fruitfly Monitoring Platform")
    title_run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = 1
    subtitle_run = subtitle.add_run("Prepared from the current implementation of the Fruitfly Monitoring Platform codebase")
    subtitle_run.italic = True

    date_para = document.add_paragraph()
    date_para.alignment = 1
    date_para.add_run("Date: 2026-04-03")

    document.add_paragraph()

    summary_table = document.add_table(rows=4, cols=2)
    summary_table.style = "Table Grid"
    summary_rows = [
        ("Project", "Fruitfly Monitoring Platform"),
        ("Frontend", "React, TypeScript, Vite"),
        ("Backend", "Node.js, Express, MySQL"),
        ("Report Scope", "Architecture, methodology, data flow, security, reporting, and implementation observations"),
    ]
    for (label, value), row in zip(summary_rows, summary_table.rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "D9EAD3")
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    document.add_paragraph()
    intro = document.add_paragraph()
    intro.alignment = 3
    intro.add_run(
        "This document presents the technical structure and operating methodology of the Fruitfly Monitoring Platform as implemented in the reviewed codebase."
    )

    document.add_section(WD_SECTION_START.NEW_PAGE)


def add_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.alignment = 1
        header.text = "Fruitfly Monitoring Platform - Technical Report"

        footer = section.footer.paragraphs[0]
        footer.alignment = 1
        footer.text = "Fruitfly Monitoring Platform"


def build_document() -> None:
    lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()
    document = Document()
    style_document(document)
    add_cover(document)

    in_numbered_list = False
    in_bullet_list = False

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            in_numbered_list = False
            in_bullet_list = False
            continue

        if stripped.startswith("# "):
            continue

        if re.match(r"^##\s+", stripped):
            para = document.add_paragraph(style="Heading 1")
            add_inline_runs(para, re.sub(r"^##\s+", "", stripped))
            in_numbered_list = False
            in_bullet_list = False
            continue

        if re.match(r"^###\s+", stripped):
            para = document.add_paragraph(style="Heading 2")
            add_inline_runs(para, re.sub(r"^###\s+", "", stripped))
            in_numbered_list = False
            in_bullet_list = False
            continue

        if re.match(r"^####\s+", stripped):
            para = document.add_paragraph(style="Heading 3")
            add_inline_runs(para, re.sub(r"^####\s+", "", stripped))
            in_numbered_list = False
            in_bullet_list = False
            continue

        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            para = document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.first_line_indent = Inches(-0.2)
            add_inline_runs(para, f"{numbered_match.group(1)}. {numbered_match.group(2)}")
            in_numbered_list = True
            in_bullet_list = False
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            para = document.add_paragraph(style="List Bullet")
            add_inline_runs(para, bullet_match.group(1))
            in_bullet_list = True
            in_numbered_list = False
            continue

        para = document.add_paragraph()
        para.alignment = 3
        add_inline_runs(para, stripped)
        in_numbered_list = False
        in_bullet_list = False

    add_header_footer(document)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
