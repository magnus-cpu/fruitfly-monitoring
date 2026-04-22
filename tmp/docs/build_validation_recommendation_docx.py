from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/magnus/Documents/Github/fruitfly-monitoring")
DOCX_OUT = ROOT / "docs" / "post_test_validation_recommended_scores.docx"


SECTIONS = [
    (
        "1. Proposed Intelligent Fruit Fly Trap - Functional Requirements",
        "These scores are inferred from the broader project description plus the available software evidence. Hardware-side items have lower confidence because this repository is primarily the monitoring platform, not the embedded trap firmware.",
        [
            ("1", "Fruit fly detection and counting", "4", "Medium", "The system clearly stores and visualizes fruit fly counts, but this repo does not fully prove the physical counting pipeline inside the trap."),
            ("2", "Image acquisition", "4", "Medium", "The backend accepts and stores base64 image captures, which supports image acquisition, though camera-side implementation is not shown here."),
            ("3", "Data transmission", "4", "Medium", "The platform exposes ingestion APIs for counts, images, telemetry, and environmental data, indicating successful upstream transmission is expected."),
            ("4", "Edge inference", "3", "Low", "There is no direct embedded or ML inference implementation in this repo, so edge inference is only partially evidenced."),
            ("5", "Power monitoring and management", "4", "Medium", "Telemetry includes voltage, current, and power, which supports power monitoring. Active power management logic is not visible here."),
            ("6", "Data storage", "5", "High", "The MySQL schema and controllers strongly support structured storage of readings, counts, images, telemetry, reports, and users."),
            ("7", "Remote monitoring and visualization", "5", "High", "This is one of the strongest proven features in the codebase through dashboards, maps, images, telemetry, and reporting."),
        ],
    ),
    (
        "2. Proposed Intelligent Fruit Fly Trap - Non-Functional Requirements",
        "These non-functional scores should be treated as recommendations rather than conclusive lab validation values. Accuracy, latency, and energy efficiency especially need hardware or field-test measurements.",
        [
            ("1", "Accuracy", "3", "Low", "The repo does not contain measurable validation data for detection accuracy or count accuracy."),
            ("2", "Latency", "3", "Low", "The web platform looks responsive, but end-to-end device-to-cloud latency is not benchmarked in this repo."),
            ("3", "Energy efficiency", "3", "Low", "Telemetry supports observation of power data, but no optimization or measured efficiency results are shown."),
            ("4", "Reliability", "4", "Medium", "The system includes structured storage, validation, and monitoring features, but uptime or failure-rate evidence is limited."),
            ("5", "Scalability", "3", "Medium", "The architecture can scale to more devices conceptually, but no load or scaling validation is demonstrated."),
            ("6", "Robustness", "3", "Low", "Some validation and fallback logic exists, but there is little evidence of rugged field robustness testing."),
            ("7", "Maintainability", "4", "Medium", "The project is modular, with separated controllers, routes, pages, and schema design, which supports maintainability."),
        ],
    ),
    (
        "3. Remote Monitoring and Visualization - Functional Requirements",
        "This section is the best-supported by the repository because the available code directly implements most of these functions.",
        [
            ("1", "Data reception", "4", "High", "The backend provides working ingestion endpoints for environmental data, counts, images, and telemetry."),
            ("2", "Data storage and management", "5", "High", "The relational database schema and API controllers strongly support storage, retrieval, filtering, and ownership control."),
            ("3", "Visualization dashboard", "5", "High", "The dashboard aggregates device status, map data, sensor metrics, image previews, and telemetry previews."),
            ("4", "Image visualization", "5", "High", "The fruitfly image page provides filters, viewer mode, download, and analysis workflow."),
            ("5", "Alerts and notifications", "2", "High", "This functionality is not strongly implemented in the current repo. There are status views, but no clear alerting/notification engine."),
            ("6", "User and access management", "4", "High", "JWT authentication, protected routes, user roles, and read-only viewer restrictions are implemented."),
            ("7", "Data export and reporting", "5", "High", "The reports module generates summary or analytics outputs and supports PDF, CSV, and JSON downloads."),
        ],
    ),
    (
        "4. Remote Monitoring and Visualization - Non-Functional Requirements",
        "These ratings are based on implemented software qualities visible in the codebase, not on formal performance testing or penetration testing.",
        [
            ("1", "Usability", "4", "Medium", "The UI is organized into clear pages for dashboard, sensors, telemetry, images, reports, and profile access."),
            ("2", "Performance", "4", "Medium", "The system appears lightweight and practical for moderate usage, though no benchmarks are included."),
            ("3", "Scalability", "3", "Medium", "The architecture is extensible, but evidence of large-scale deployment or optimization is limited."),
            ("4", "Reliability and availability", "4", "Medium", "Structured APIs, health checks, and clear route separation support reliability, but no uptime evidence is shown."),
            ("5", "Security and privacy", "4", "High", "JWT auth, role restrictions, rate limiting, Helmet, and CORS are present, though device-level auth is still limited."),
            ("6", "Interoperability", "3", "Low", "The system uses standard web APIs and JSON, but broader interoperability integrations are not demonstrated."),
            ("7", "Maintainability", "4", "Medium", "The modular code structure supports maintenance, though documentation and automated tests can be improved."),
            ("8", "Cost efficiency", "4", "Medium", "The stack is practical and cost-conscious for a monitoring platform, but no formal cost analysis is included."),
        ],
    ),
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_margins(section) -> None:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    set_margins(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(20)
    title.font.bold = True

    for name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.text = "Post-Test Validation - Recommended Scores"

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Evidence-based recommendations derived from the current Fruitfly Monitoring Platform codebase"


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Post-Test Validation Recommended Scores")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Project: Fruitfly Monitoring Platform").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared: 2026-04-03")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "This document converts the validation questionnaire into an evidence-based recommendation sheet. "
        "The suggested scores are derived from the available codebase and should be used as guidance for filling the validation form. "
        "Items that depend on hardware testing, field measurements, or embedded inference have lower confidence than items directly implemented in the web platform."
    )

    key = doc.add_table(rows=1, cols=2)
    key.style = "Table Grid"
    key.rows[0].cells[0].text = "Scoring key"
    key.rows[0].cells[1].text = "1 = Very low, 2 = Low, 3 = Average, 4 = High, 5 = Very High"
    shade(key.rows[0].cells[0], "D9EAD3")
    key.rows[0].cells[0].paragraphs[0].runs[0].bold = True

    key2 = doc.add_table(rows=1, cols=2)
    key2.style = "Table Grid"
    key2.rows[0].cells[0].text = "Important note"
    key2.rows[0].cells[1].text = "Sections on trap hardware should ideally be confirmed with prototype, field, or laboratory evidence in addition to this software repo."
    shade(key2.rows[0].cells[0], "FCE5CD")
    key2.rows[0].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_section_table(doc: Document, title: str, intro: str, rows: list[tuple[str, str, str, str, str]]) -> None:
    h = doc.add_paragraph(style="Heading 1")
    h.add_run(title)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(intro)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["SN", "Criterion", "Suggested Score", "Confidence", "Reason / Evidence Basis"]
    colors = ["D9EAD3"] * 5
    for cell, text, fill in zip(hdr, headers, colors):
        set_cell_text(cell, text, bold=True, size=10)
        shade(cell, fill)

    for sn, criterion, score, confidence, reason in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], sn)
        set_cell_text(cells[1], criterion)
        set_cell_text(cells[2], score, bold=True)
        set_cell_text(cells[3], confidence)
        set_cell_text(cells[4], reason)

    table.columns[0].width = Inches(0.45)
    table.columns[1].width = Inches(2.2)
    table.columns[2].width = Inches(1.0)
    table.columns[3].width = Inches(1.0)
    table.columns[4].width = Inches(5.0)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_cover(doc)

    for index, (title, intro, rows) in enumerate(SECTIONS):
        if index > 0:
            doc.add_page_break()
        add_section_table(doc, title, intro, rows)

    doc.add_page_break()
    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Overall Guidance for Filling the Original Form")

    guidance = [
        "Use the suggested score as a recommendation, not as an unquestionable final mark.",
        "For Sections 3 and 4, the recommended scores are strongly supported by the available web-platform code.",
        "For Sections 1 and 2, confirm hardware-side answers using prototype tests, embedded inference results, transmission logs, and field observations if available.",
        "The weakest area in the current software evidence is Alerts and notifications, which is why it receives a lower score.",
        "If your supervisor or validators expect conservative scoring, keep low-confidence items at 3 unless independent test evidence justifies 4 or 5.",
    ]
    for item in guidance:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_header_footer(doc)
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


if __name__ == "__main__":
    build()
