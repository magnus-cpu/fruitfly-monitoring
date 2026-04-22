from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/magnus/Documents/Github/fruitfly-monitoring")
DOCX_OUT = ROOT / "docs" / "web_validation_sections_3_4.docx"


SECTION_3 = [
    ("1", "Data reception", "4", "The backend receives environmental data, counts, images, and telemetry through working APIs."),
    ("2", "Data storage and management", "5", "The platform has a structured MySQL schema with controllers for retrieval, filtering, and ownership control."),
    ("3", "Visualization dashboard", "5", "The dashboard presents sensor status, map data, metrics, image preview, and telemetry preview."),
    ("4", "Image visualization", "5", "The system includes image listing, filtering, preview, analysis status, and download features."),
    ("5", "Alerts and notifications", "2", "This is not strongly implemented in the current repository."),
    ("6", "User and access management", "4", "JWT authentication, protected routes, user roles, and viewer restrictions are implemented."),
    ("7", "Data export and reporting", "5", "The reports module supports summary and analytics generation with PDF, CSV, and JSON downloads."),
]


SECTION_4 = [
    ("1", "Usability", "4", "The UI is organized into clear pages for dashboard, sensors, telemetry, images, reports, and profile access."),
    ("2", "Performance", "4", "The platform appears lightweight and practical for moderate usage, although no benchmark data is included."),
    ("3", "Scalability", "3", "The architecture is extensible, but large-scale validation is not demonstrated in this repository."),
    ("4", "Reliability and availability", "4", "The system includes structured APIs, route separation, and a health check, but no uptime evidence is shown."),
    ("5", "Security and privacy", "4", "JWT auth, role restrictions, Helmet, rate limiting, and CORS are implemented."),
    ("6", "Interoperability", "3", "The system uses standard JSON APIs, but broader integration evidence is limited."),
    ("7", "Maintainability", "4", "The codebase is modular, with separated pages, routes, controllers, and schema definitions."),
    ("8", "Cost efficiency", "4", "The chosen stack is practical and cost-conscious for a monitoring platform, though no formal cost study is included."),
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(20)
    title.font.bold = True

    heading = doc.styles["Heading 1"]
    heading.font.name = "Times New Roman"
    heading.font.size = Pt(15)
    heading.font.bold = True


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.text = "Validation Response For Web Platform - Sections 3 And 4"

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Fruitfly Monitoring Platform"


def add_table(doc: Document, title: str, intro: str, rows: list[tuple[str, str, str, str]]) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run(title)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(intro)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["SN", "Criterion", "Score", "Short Justification"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, size=10)
        shade(cell, "D9EAD3")

    for sn, criterion, score, reason in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], sn)
        set_cell_text(cells[1], criterion)
        set_cell_text(cells[2], score, bold=True)
        set_cell_text(cells[3], reason)

    table.columns[0].width = Inches(0.45)
    table.columns[1].width = Inches(2.4)
    table.columns[2].width = Inches(0.8)
    table.columns[3].width = Inches(6.0)


def build() -> None:
    doc = Document()
    setup_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Validation Response For Web Platform")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Sections 3 and 4 Only").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared: 2026-04-03")

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "This submission-ready document contains recommended answers for Section 3 "
        "(remote monitoring and visualization functional requirements) and Section 4 "
        "(remote monitoring and visualization non-functional requirements) based on the "
        "available Fruitfly Monitoring Platform codebase."
    )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note_run = note.add_run(
        "Scoring key: 1 = Very low, 2 = Low, 3 = Average, 4 = High, 5 = Very High."
    )
    note_run.bold = True

    doc.add_page_break()

    add_table(
        doc,
        "Section 3. Remote Monitoring And Visualization - Functional Requirements",
        "These scores are strongly supported by the implemented web platform features.",
        SECTION_3,
    )

    doc.add_page_break()

    add_table(
        doc,
        "Section 4. Remote Monitoring And Visualization - Non-Functional Requirements",
        "These scores are based on visible software qualities in the repository and should be treated as evidence-based recommendations.",
        SECTION_4,
    )

    add_header_footer(doc)
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


if __name__ == "__main__":
    build()
